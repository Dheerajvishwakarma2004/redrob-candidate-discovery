"""
JD Document Loader.
Supports .docx, .md, .txt.
Returns plain text string for all formats.
"""

import re
from pathlib import Path
from src.utils.logger import get_logger

logger = get_logger(__name__)

SUPPORTED = {".docx", ".md", ".txt", ".text"}


class DocumentLoader:

    def load(self, path: str) -> str:
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(f"JD file not found: {p.absolute()}")
        suffix = p.suffix.lower()
        if suffix not in SUPPORTED:
            raise ValueError(f"Unsupported format: {suffix}. Supported: {SUPPORTED}")
        logger.info(f"Loading JD: {p} [{suffix}]")
        if suffix == ".docx":
            text = self._load_docx(p)
        elif suffix == ".md":
            text = self._load_markdown(p)
        else:
            text = self._load_text(p)
        text = self._clean(text)
        if not text.strip():
            raise ValueError(f"JD is empty after extraction: {p}")
        logger.info(f"JD loaded: {len(text)} chars, ~{len(text.split())} words")
        return text

    def _load_docx(self, p: Path) -> str:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")
        doc = Document(str(p))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)

    def _load_text(self, p: Path) -> str:
        for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                return p.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Cannot decode: {p}")

    def _load_markdown(self, p: Path) -> str:
        raw = self._load_text(p)
        text = re.sub(r"^#{1,6}\s+", "", raw, flags=re.MULTILINE)
        text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
        text = re.sub(r"`{3}.*?`{3}", "", text, flags=re.DOTALL)
        text = re.sub(r"`(.+?)`", r"\1", text)
        text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
        text = re.sub(r"^[\*\-\+]\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"^\d+\.\s+", "", text, flags=re.MULTILINE)
        return text

    def _clean(self, text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[^\x09\x0A\x20-\x7E\x80-\xFF]", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]{2,}", " ", text)
        return text.strip()