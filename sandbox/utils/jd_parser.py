# sandbox/utils/jd_parser.py
import re
import os
from pathlib import Path
from typing import Dict, List, Tuple
import numpy as np
from sentence_transformers import SentenceTransformer

CONCEPT_DESCRIPTIONS = {
    "retrieval": (
        "search retrieval information retrieval vector search BM25 FAISS "
        "Elasticsearch semantic search query understanding document retrieval "
        "search engine dense retrieval sparse retrieval nearest neighbor"
    ),
    "ranking": (
        "ranking learning to rank relevance scoring NDCG MAP ranking model "
        "reranking relevance modeling candidate ranking result ranking "
        "pointwise pairwise listwise LambdaMART"
    ),
    "recommendation": (
        "recommendation systems recommender collaborative filtering "
        "personalization matching engine candidate discovery job matching "
        "content-based filtering matrix factorization similar profiles"
    ),
    "ml_engineering": (
        "machine learning MLOps feature engineering model deployment "
        "training pipeline production ML model serving experiment tracking "
        "A/B testing airflow spark data pipeline scikit-learn xgboost"
    ),
    "llm": (
        "large language models LLM RAG fine-tuning embeddings transformers "
        "prompt engineering NLP BERT GPT huggingface sentence embeddings "
        "text classification natural language processing"
    ),
}

EXPERIENCE_PATTERNS = [
    r"(\d+)\s*[-–to]+\s*(\d+)\s*years?",
    r"(\d+)\+\s*years?",
    r"(?:minimum|at least|min\.?)\s*(\d+)\s*years?",
]

KNOWN_TITLES = [
    "search engineer", "retrieval engineer", "recommendation systems engineer",
    "recommendation engineer", "ranking engineer", "relevance engineer",
    "ml engineer", "machine learning engineer", "ai engineer",
    "nlp engineer", "data scientist", "applied scientist",
    "research engineer", "software engineer", "backend engineer",
    "data engineer", "mlops engineer",
]

def get_embedding_model() -> SentenceTransformer:
    """Helper to load embedding model locally if possible."""
    local_paths = [
        Path(__file__).parent.parent.parent / "models" / "bge-small-en-v1.5",
        Path(__file__).parent.parent / "models" / "bge-small-en-v1.5",
        Path("models/bge-small-en-v1.5")
    ]
    for p in local_paths:
        if p.exists() and (p / "pytorch_model.bin").exists():
            return SentenceTransformer(str(p))
    raise FileNotFoundError("Local model directory models/bge-small-en-v1.5 not found. Run download_artifacts.py first.")

class DocumentLoader:
    def load_text(self, text: str) -> str:
        return self._clean(text)

    def load_file(self, file_content: bytes, file_name: str) -> str:
        suffix = Path(file_name).suffix.lower()
        if suffix == ".docx":
            return self._load_docx(file_content)
        elif suffix == ".md":
            return self._load_markdown(file_content.decode("utf-8", errors="ignore"))
        else:
            return self._clean(file_content.decode("utf-8", errors="ignore"))

    def _load_docx(self, content: bytes) -> str:
        import io
        from docx import Document
        doc = Document(io.BytesIO(content))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return self._clean("\n".join(parts))

    def _load_markdown(self, raw: str) -> str:
        text = re.sub(r"^#{1,6}\s+", "", raw, flags=re.MULTILINE)
        text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
        text = re.sub(r"`{3}.*?`{3}", "", text, flags=re.DOTALL)
        text = re.sub(r"`(.+?)`", r"\1", text)
        text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
        text = re.sub(r"^[\*\-\+]\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"^\d+\.\s+", "", text, flags=re.MULTILINE)
        return self._clean(text)

    def _clean(self, text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[^\x09\x0A\x20-\x7E\x80-\xFF]", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]{2,}", " ", text)
        return text.strip()

class JDExtractor:
    def __init__(self, model: SentenceTransformer = None):
        self._model = model

    @property
    def model(self) -> SentenceTransformer:
        if self._model is None:
            self._model = get_embedding_model()
        return self._model

    def extract(self, jd_text: str) -> Dict:
        clean = jd_text.lower()
        # Direct encode without instruction to match production cache embedding logic
        jd_embedding = self.model.encode(jd_text, normalize_embeddings=True, show_progress_bar=False)
        jd_embedding = jd_embedding.astype(np.float32)

        # Concept weights
        concept_weights = {}
        for name, desc in CONCEPT_DESCRIPTIONS.items():
            desc_emb = self.model.encode(desc, normalize_embeddings=True, show_progress_bar=False)
            concept_weights[name] = float(max(np.dot(jd_embedding, desc_emb.astype(np.float32)), 0.0))

        # Check threshold
        threshold = 0.28
        if max(concept_weights.values()) < threshold:
            fallback = {
                "retrieval": 0.30,
                "ranking": 0.25,
                "recommendation": 0.20,
                "ml_engineering": 0.20,
                "llm": 0.05
            }
            concept_weights = fallback
        else:
            total = sum(concept_weights.values()) or 1.0
            concept_weights = {k: v / total for k, v in concept_weights.items()}

        exp_min, exp_max = self._extract_experience(clean)
        titles = [t for t in KNOWN_TITLES if t in clean]
        primary_domain = max(concept_weights, key=concept_weights.get)

        return {
            "raw_text": jd_text,
            "clean_text": clean,
            "jd_embedding": jd_embedding,
            "concept_weights": concept_weights,
            "primary_domain": primary_domain,
            "experience_min": exp_min,
            "experience_max": exp_max,
            "required_titles": titles,
        }

    def _extract_experience(self, text: str) -> Tuple[float, float]:
        for pattern in EXPERIENCE_PATTERNS:
            m = re.search(pattern, text)
            if m:
                groups = m.groups()
                if len(groups) == 2 and groups[1]:
                    return float(groups[0]), float(groups[1])
                elif groups[0]:
                    mn = float(groups[0])
                    return mn, mn + 4.0
        return 5.0, 9.0
